import os
import uuid
import json
import requests
from datetime import datetime
from docx import Document
from openai import OpenAI
import base64
from collections import defaultdict


def generate_family_report(data, images, api_key):
    """
    Genereert een familieverslag met behulp van de OpenRouter API.

    Parameters:
    - data: Een dictionary met de gegevens voor het verslag
    - images: Een lijst met geüploade afbeeldingspaden
    - api_key: De OpenRouter API-sleutel

    Returns:
    - Een tuple (verslag_tekst, bestandsnaam_docx)
    """
    # Converteer de afbeeldingen naar base64 voor de API
    image_descriptions = []
    for img_path in images:
        if os.path.exists(img_path):
            with open(img_path, 'rb') as img_file:
                file_extension = os.path.splitext(img_path)[1][1:].lower()
                img_base64 = base64.b64encode(img_file.read()).decode('utf-8')
                image_descriptions.append(f"![Afbeelding](data:image/{file_extension};base64,{img_base64})")

    # Bereid de input voor de AI voor
    prompt = f"""
    Schrijf een volledig familiestamboom onderzoeksverslag op basis van de volgende informatie:

    ONDERZOEKSVRAAG: {data.get('onderzoeksvraag', '')}
    DEELVRAGEN: {data.get('deelvragen', '')}

    INTERVIEWS:
    {data.get('interviews', '')}

    BRONNEN:
    {data.get('bronnen', '')}

    DOCUMENTEN EN BOEKEN:
    {data.get('documenten', '')}

    EXTRA NOTITIES:
    {data.get('notities', '')}

    Het verslag moet het volgende bevatten:
    1. Een inleiding met de onderzoeksvraag en deelvragen
    2. Een beschrijving van hoe het onderzoek is aangepakt
    3. De resultaten van het onderzoek, inclusief wat er is ontdekt
    4. Een conclusie die antwoord geeft op de onderzoeksvraag

    Schrijf dit als een samenhangend verhaal over de familiegeschiedenis, met een duidelijke structuur en hoofdstukken.
    Noem specifieke feiten, plaatsen, jaartallen, en personen die in de input worden genoemd.
    """

    # OpenRouter API aanroepen
    try:
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )

        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "familiegeschiedenisproject.com",
                "X-Title": "Familiegeschiedenis Project",
            },
            model="deepseek/deepseek-r1-zero:free",
            messages=[
                {"role": "system",
                 "content": "Je bent een expert in familiestamboomonderzoek en geschiedschrijving. Je schrijft een uitgebreide onderzoek met alle data die je hebt gekregen. Je zoekt op het internet naar de gegeven bronnen en documenten."},
                {"role": "user", "content": prompt}
            ]
        )

        report_text = completion.choices[0].message.content

        # Maak een Word-document
        doc_path = create_word_document(data, report_text, images)

        return report_text, doc_path

    except Exception as e:
        print(f"Fout bij het genereren van het verslag: {e}")
        return f"Er is een fout opgetreden bij het genereren van het verslag: {str(e)}", None


def create_word_document(data, report_text, images):
    """Maakt een Word-document met het gegenereerde verslag en afbeeldingen"""
    doc = Document()

    # Titel
    doc.add_heading('Familiestamboom Onderzoeksverslag', 0)

    # Datum
    doc.add_paragraph(f"Gegenereerd op: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

    # Onderzoeksvraag
    doc.add_heading('Onderzoeksvraag', 1)
    doc.add_paragraph(data.get('onderzoeksvraag', ''))

    # Verslag tekst
    doc.add_heading('Verslag', 1)
    for paragraph in report_text.split('\n\n'):
        if paragraph.strip():
            # Controleer op kopjes
            if paragraph.strip().startswith('#'):
                level = paragraph.count('#')
                doc.add_heading(paragraph.strip('#').strip(), level)
            else:
                doc.add_paragraph(paragraph)

    # Afbeeldingen toevoegen
    if images:
        doc.add_heading('Bijgevoegde afbeeldingen', 1)
        for img_path in images:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=4000000)  # Breedte in EMU (ca. 10cm)
                    doc.add_paragraph(f"Afbeelding: {os.path.basename(img_path)}")
                except Exception as e:
                    doc.add_paragraph(f"Kon afbeelding {os.path.basename(img_path)} niet toevoegen: {str(e)}")

    # Sla het bestand op
    unique_id = uuid.uuid4().hex[:8]
    filename = f"familiestamboom_verslag_{unique_id}.docx"
    file_path = os.path.join('app/static/uploads', filename)
    doc.save(file_path)

    return filename


def generate_family_tree_data(data):
    """
    Genereert gedetailleerde familierelaties uit de verstrekte gegevens
    voor visualisatie in een familiestamboom met behulp van een language model.
    """
    # Combineer alle relevante gegevens
    interviews = data.get('interviews', '')
    bronnen = data.get('bronnen', '')
    documenten = data.get('documenten', '')
    onderzoeksvraag = data.get('onderzoeksvraag', '')
    deelvragen = data.get('deelvragen', '')
    notities = data.get('notities', '')

    all_text = f"Onderzoeksvraag: {onderzoeksvraag}\n\nDeelvragen: {deelvragen}\n\n" \
               f"Interviews: {interviews}\n\nBronnen: {bronnen}\n\n" \
               f"Documenten: {documenten}\n\nNotities: {notities}"

    # Gebruik het LLM om de familierelaties te analyseren
    family_data = extract_family_relationships_with_llm(all_text)

    return json.dumps(family_data)


def extract_family_relationships_with_llm(text):
    """
    Gebruikt een language model om familierelaties te identificeren en te structureren.
    """
    # API key die al in de routes.py wordt gebruikt
    api_key = "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"

    # Maak verzoek naar het language model
    try:
        response = make_llm_request(text, api_key)
        family_tree = process_llm_response(response)
        return family_tree
    except Exception as e:
        print(f"Fout bij het verwerken van familierelaties: {e}")
        return fallback_family_tree()


def make_llm_request(text, api_key):
    """
    Stuur verzoek naar het language model om familierelaties te analyseren,
    gebruikmakend van de OpenRouter API.
    """
    try:
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )

        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "familiegeschiedenisproject.com",
                "X-Title": "Familiegeschiedenis Project",
            },
            model="deepseek/deepseek-r1-zero:free",
            messages=[
                {"role": "system",
                 "content": "Je bent een expert in het analyseren van familierelaties en stambomen. Je kunt personen en hun onderlinge verbanden identificeren uit teksten."},
                {"role": "user", "content": f"""
                Analyseer de volgende tekst en identificeer alle genoemde personen en hun familierelaties:
                
                {text}
                
                Geef het resultaat terug in het volgende JSON-formaat:
                {{
                    "personen": [
                        {{
                            "id": "1",
                            "naam": "Volledige naam",
                            "geboorte": "geboortedatum (indien bekend)",
                            "overlijden": "overlijdensdatum (indien bekend)"
                        }}
                    ],
                    "relaties": [
                        {{
                            "van": "id van persoon 1",
                            "naar": "id van persoon 2",
                            "type": "relatietype (ouder, kind, partner, etc.)"
                        }}
                    ]
                }}
                
                Als de informatie niet beschikbaar is, gebruik dan "onbekend".
                """}
            ]
        )

        # Extract content from the response
        content = completion.choices[0].message.content

        # Try to extract the JSON part of the response
        import re
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content)
        if json_match:
            json_str = json_match.group(1).strip()
        else:
            json_str = content

        # Clean up the JSON string to make sure it's valid
        json_str = json_str.replace('```json', '').replace('```', '')
        return json.loads(json_str)

    except Exception as e:
        print(f"Fout bij LLM aanroep: {e}")
        if 'completion' in locals():
            print(f"Response: {completion}")
        raise e


def process_llm_response(response):
    """
    Verwerk de response van het language model tot een stamboom-structuur
    die geschikt is voor visualisatie.
    """
    # Creëer een boom-structuur op basis van de LLM-resultaten
    tree = {"name": "Familiestamboom", "children": []}

    # Als de response geen geldige personen of relaties heeft, gebruik fallback
    if 'personen' not in response or not response['personen']:
        return fallback_family_tree()

    # Maak een mapping van persoon ID naar persoonsobject
    persons = {}
    for person in response['personen']:
        person_id = person.get('id', str(len(persons) + 1))
        person_name = person.get('naam', 'Onbekend')
        birth = person.get('geboorte', '')
        death = person.get('overlijden', '')

        # Creëer knooppunt voor deze persoon
        display_name = person_name
        if birth and birth != "onbekend":
            display_name += f" ({birth}"
            if death and death != "onbekend":
                display_name += f" - {death}"
            display_name += ")"

        persons[person_id] = {
            "name": display_name,
            "full_data": person,
            "children": []
        }

    # Verwerk relaties
    if 'relaties' in response:
        # Maak een mapping van ouder-kind relaties
        family_graph = defaultdict(list)

        for relation in response['relaties']:
            source_id = relation.get('van')
            target_id = relation.get('naar')
            rel_type = relation.get('type', '').lower()

            # Voeg alleen relaties toe als beide personen bestaan
            if source_id in persons and target_id in persons:
                if rel_type in ['ouder', 'vader', 'moeder', 'parent']:
                    family_graph[source_id].append(target_id)
                elif rel_type in ['kind', 'zoon', 'dochter', 'child']:
                    family_graph[target_id].append(source_id)

        # Identificeer de roots (personen zonder ouders)
        all_persons = set(persons.keys())
        all_children = set()

        for parent, children in family_graph.items():
            for child in children:
                all_children.add(child)

        roots = all_persons - all_children

        # Als er geen roots zijn, beschouw dan alle personen als roots
        if not roots:
            roots = all_persons

        # Bouw de boomstructuur op
        def build_subtree(person_id, visited=None):
            if visited is None:
                visited = set()

            if person_id in visited:
                return None  # Voorkom oneindige recursie

            visited.add(person_id)

            node = dict(persons[person_id])
            children = []

            if person_id in family_graph:
                for child_id in family_graph[person_id]:
                    child_node = build_subtree(child_id, visited.copy())
                    if child_node:
                        children.append(child_node)

            if children:
                node["children"] = children

            return node

        # Bouw de boom op met de geïdentificeerde roots
        for root in roots:
            subtree = build_subtree(root)
            if subtree:
                tree["children"].append(subtree)

    # Als er geen kinderen zijn, voeg alle personen toe als directe kinderen van de root
    if not tree["children"]:
        for person_id, person in persons.items():
            tree["children"].append(person)

    return tree


def fallback_family_tree():
    """
    Creëer een eenvoudige fallback-stamboom als er geen relaties konden worden geëxtraheerd.
    """
    return {
        "name": "Familiestamboom",
        "children": [
            {
                "name": "Geen familierelaties gevonden",
                "children": [
                    {"name": "Voeg meer details toe over familieleden in de tekstvelden."},
                    {"name": "Vermeld namen, geboortedata, en relaties tussen personen."}
                ]
            }
        ]
    }


def extract_names_from_text(text):
    """Zeer eenvoudige functie om mogelijke namen uit tekst te halen"""
    # Dit is een zeer eenvoudige implementatie die alleen werkt voor namen tussen dubbele quotes
    import re
    names = re.findall(r'"([^"]*)"', text)

    # Voeg enkele standaard familienamen toe als er geen werden gevonden
    if not names:
        names = ["No names Given", "No names Given", "No names Given", "No names Given", "No names Given"]

    return names[:10]  # Limiteer tot 10 namen
